from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "projekt-psk.docx"
GODLO = ROOT / "docs" / "godlo.jpg"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_paragraph(doc, text="", style=None, align=None, bold=False):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return paragraph


def add_equation(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.right_indent = Cm(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    for line_index, line in enumerate(text.strip("\n").splitlines()):
        if line_index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    paragraph._p.get_or_add_pPr().append(shading)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(6, 95, 70)),
        ("Heading 2", 14, RGBColor(20, 83, 45)),
        ("Heading 3", 12, RGBColor(20, 83, 45)),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for text in [
        "Collegium Witelona Uczelnia Państwowa w Legnicy",
        "Wydział Nauk Technicznych i Ekonomicznych",
        "Kierunek: Informatyka",
    ]:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc)
    add_paragraph(doc)

    if GODLO.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(GODLO), width=Cm(5))
    else:
        for _ in range(4):
            add_paragraph(doc)

    add_paragraph(doc)
    title = add_paragraph(
        doc,
        "Projekt do kursu Podstawy Symulacji Komputerowej''",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    title.runs[0].font.size = Pt(14)

    for _ in range(5):
        add_paragraph(doc)

    author = add_paragraph(doc, "Autor", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    author.runs[0].font.size = Pt(14)
    add_paragraph(doc, "Gabriela Grabarska, nr indeksu: 43840", align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(4):
        add_paragraph(doc)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    cell = table.cell(0, 0)
    cell.text = "Prowadzący przedmiot\ndr inż. Ryszard Rębowski"
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    for _ in range(5):
        add_paragraph(doc)

    add_paragraph(doc, "Legnica, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def add_probability_table(doc):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    values = [
        ["", "Y = 3", "Y = 2", "Y = 1"],
        ["X = 1", "0,1", "0,2", "0,3"],
        ["X = 2", "0,1", "0,1", "0,2"],
    ]
    for r, row in enumerate(values):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0 or c == 0))
            if r == 0 or c == 0:
                set_cell_shading(table.cell(r, c), "324E78")
                set_cell_text(table.cell(r, c), value, bold=True, color=(255, 255, 255))
    return table


def build():
    doc = Document()
    setup_styles(doc)
    add_title_page(doc)

    add_paragraph(doc, "Spis treści", style="Heading 1")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    doc.add_heading("Wprowadzenie", level=1)
    add_paragraph(
        doc,
        "Celem projektu jest opracowanie trzech tematów z list zadań z kursu Podstawy "
        "Symulacji Komputerowej. Projekt obejmuje część merytoryczną, implementację "
        "algorytmów oraz prezentację wyników końcowych w aplikacji komputerowej.",
    )
    add_paragraph(
        doc,
        "W pracy przedstawiono podstawowe zagadnienia związane z symulacją komputerową, "
        "generowaniem liczb pseudolosowych oraz wykorzystaniem tych liczb do otrzymywania "
        "rozkładów prawdopodobieństwa. Dla każdego wybranego zadania opisano podstawy "
        "merytoryczne, sposób implementacji oraz otrzymane wyniki.",
    )
    add_paragraph(
        doc,
        "Do realizacji projektu przygotowano aplikację webową. Część obliczeniowa została "
        "wydzielona do backendu w języku Python, natomiast część prezentacyjna została "
        "wykonana w React.",
    )
    add_bullets(
        doc,
        [
            "Lista 0, zadanie 9 - analiza rozkładu łącznego wektora losowego.",
            "Lista 1, zadanie 8 - generowanie rozkładu Poissona.",
            "Lista 2, zadanie 2 - generowanie rozkładu ciągłego metodą odwracania dystrybuanty.",
        ],
    )

    doc.add_heading("Generator LCG Lehmera", level=1)
    add_paragraph(
        doc,
        "W zadaniach z Listy 1 i Listy 2 wykorzystano generator kongruencyjny Lehmera "
        "jako źródło liczb pseudolosowych o rozkładzie jednostajnym na przedziale [0,1).",
    )
    add_equation(doc, "X_j = aX_{j-1} mod m")
    add_equation(doc, "U_j = X_j / m")
    add_equation(doc, "L = round(log_2(k) + 2),   m = 2^L")
    add_bullets(
        doc,
        [
            "L > 4",
            "k >= 100",
            "a mod 8 należy do {3, 5}",
            "X_0 jest liczbą nieparzystą",
        ],
    )
    add_equation(doc, "k = 536870912,   a = 1103515245,   X_0 = 12345,   n = 1200")
    add_equation(doc, "L = 31,   m = 2^31,   okres = 2^29 = 536870912")

    doc.add_heading("Obserwacje i wnioski dotyczące doboru parametrów", level=2)
    add_paragraph(
        doc,
        "Początkowo zastosowano parametry k = 100, a = 101, X_0 = 3 oraz n = 900. "
        "Dla tych wartości na wykresach 2D i 3D widoczne były wyraźne linie oraz "
        "regularne układy punktów.",
    )
    add_equation(doc, "dla k = 100:   L = 9,   m = 2^9 = 512,   okres = 2^7 = 128")
    add_paragraph(
        doc,
        "Liczba generowanych punktów była znacznie większa niż okres generatora. "
        "W konsekwencji ciąg zaczynał się wielokrotnie powtarzać, a ta powtarzalność "
        "była widoczna na wykresach.",
    )
    add_paragraph(
        doc,
        "Po zgłębieniu tematu generatorów liniowych kongruencyjnych zauważono, że jakość "
        "generatora silnie zależy od doboru parametrów a, m oraz X_0. W opracowaniu Karla "
        "Entachera przedstawiono klasyczny generator używany w ANSI C.",
    )
    add_equation(doc, "LCG(2^31, 1103515245, 12345, 12345)")
    add_paragraph(
        doc,
        "W projekcie nie zastosowano dokładnie generatora ANSI C, ponieważ wzór z zadania "
        "dotyczy generatora multiplikatywnego Lehmera, bez składnika addytywnego.",
    )
    add_equation(doc, "X_j = aX_{j-1} mod m")
    add_equation(doc, "X_j = aX_{j-1} + b mod m")
    add_paragraph(
        doc,
        "Z przytoczonego przykładu wykorzystano ideę dużego modułu 2^31 oraz znanego "
        "mnożnika 1103515245. Dzięki temu ciąg nie powtarza się szybko, a punkty na "
        "wykresach są znacznie lepiej rozproszone.",
    )

    doc.add_heading("Fragment implementacji", level=2)
    add_code(
        doc,
        """
@dataclass
class LehmerGenerator:
    seed: int
    modulus: int
    multiplier: int

    def next_int(self) -> int:
        self.current_x = (self.multiplier * self.current_x) % self.modulus
        return self.current_x

    def next_float(self) -> float:
        return self.next_int() / self.modulus
""",
    )

    doc.add_heading("Drugi sposób generowania liczb pseudolosowych: metoda von Neumanna", level=1)
    add_paragraph(
        doc,
        "Metoda von Neumanna, nazywana metodą środkowych kwadratów, polega na podniesieniu "
        "poprzedniej wartości do kwadratu i wybraniu środkowych cyfr otrzymanego wyniku.",
    )
    add_numbered(
        doc,
        [
            "Pobierana jest poprzednia wartość X_{n-1}.",
            "Obliczany jest kwadrat Y_n = X_{n-1}^2.",
            "Wynik jest uzupełniany zerami tak, aby miał 2m cyfr.",
            "Ze środka zapisu liczby wybieranych jest m cyfr.",
            "Otrzymana liczba staje się nową wartością X_n.",
        ],
    )
    add_equation(doc, "X_{n-1} -> X_{n-1}^2 -> środkowe m cyfr -> X_n")
    add_equation(doc, "12^2 = 144 -> 0144 -> 14")
    add_equation(doc, "14^2 = 196 -> 0196 -> 19")
    add_equation(doc, "19^2 = 361 -> 0361 -> 36")
    add_paragraph(
        doc,
        "Zaletą tej metody jest prostota i łatwość wizualnego pokazania kolejnych kroków. "
        "Jej ograniczeniem jest możliwość szybkiego wejścia w krótki cykl albo osiągnięcia "
        "wartości 0.",
    )
    add_code(
        doc,
        """
def calculate_von_neumann(seed: int, digits: int, count: int) -> list[dict]:
    current_x = int(seed)
    results = []

    for index in range(count):
        square = current_x**2
        square_text = str(square).zfill(2 * digits)
        start = digits // 2
        middle = square_text[start : start + digits]
        value = int(middle)
        current_x = value

        if current_x == 0:
            break

    return results
""",
    )

    doc.add_heading("Lista 0, zadanie 9", level=1)
    doc.add_heading("Treść zadania", level=2)
    add_paragraph(doc, "Dana jest macierz P reprezentująca rozkład łączny wektora losowego (X,Y):")
    add_probability_table(doc)
    add_equation(doc, "X(Ω) = {1, 2},   Y(Ω) = {3, 2, 1}")
    add_bullets(
        doc,
        [
            "Obliczyć cov(X,Y).",
            "Obliczyć współczynnik korelacji ρ(X,Y).",
            "Sprawdzić niezależność stochastyczną.",
            "Sprawdzić zależność liniową.",
        ],
    )
    doc.add_heading("Podstawy merytoryczne i obliczenia", level=2)
    add_equation(doc, "p_ij = P(X = x_i, Y = y_j)")
    add_equation(doc, "P(X = x_i) = Σ_j p_ij,   P(Y = y_j) = Σ_i p_ij")
    add_equation(doc, "P(X=1)=0,6,   P(X=2)=0,4")
    add_equation(doc, "P(Y=3)=0,2,   P(Y=2)=0,3,   P(Y=1)=0,5")
    add_equation(doc, "E(X)=1·0,6+2·0,4=1,4")
    add_equation(doc, "E(Y)=3·0,2+2·0,3+1·0,5=1,7")
    add_equation(doc, "E(XY)=2,4")
    add_equation(doc, "cov(X,Y)=E(XY)-E(X)E(Y)=2,4-1,4·1,7=0,02")
    add_equation(doc, "Var(X)=0,24,   Var(Y)=0,61")
    add_equation(doc, "ρ(X,Y)=0,02 / sqrt(0,24·0,61) ≈ 0,0522")
    add_equation(doc, "P(X=1,Y=3)=0,1,   P(X=1)P(Y=3)=0,6·0,2=0,12")
    add_paragraph(doc, "Ponieważ 0,1 != 0,12, zmienne X i Y nie są stochastycznie niezależne.")
    add_paragraph(doc, "Zmienne nie są również liniowo zależne, ponieważ dla X = 1 zmienna Y może przyjmować kilka różnych wartości z dodatnim prawdopodobieństwem.")
    add_code(
        doc,
        """
expected_xy = 0.0
for row_index, x in enumerate(x_values):
    for column_index, y in enumerate(y_values):
        expected_xy += x * y * probabilities[row_index][column_index]

covariance = expected_xy - expected_x * expected_y
rho = covariance / sqrt(variance_x * variance_y)
""",
    )

    doc.add_heading("Lista 1, zadanie 8", level=1)
    doc.add_heading("Treść zadania", level=2)
    add_paragraph(doc, "Zaprogramować algorytm generowania rozkładu Poissona. Przyjąć:")
    add_equation(doc, "λ = 2,0")
    add_paragraph(doc, "W projekcie liczby jednostajne U_i generowane są za pomocą generatora LCG Lehmera.")
    doc.add_heading("Podstawy merytoryczne", level=2)
    add_equation(doc, "P(X=k)=e^{-λ} λ^k / k!,   k=0,1,2,...")
    add_equation(doc, "dla λ=2:   P(X=k)=e^{-2} 2^k / k!")
    add_equation(doc, "E(X)=λ,   Var(X)=λ")
    add_equation(doc, "F(k)=P(X≤k)=Σ_{i=0}^{k} e^{-λ} λ^i / i!")
    add_equation(doc, "F(k)≥U")
    doc.add_heading("Algorytm", level=2)
    add_numbered(
        doc,
        [
            "Wygenerować wartość U algorytmem LCG Lehmera.",
            "Ustawić p_0 = e^{-λ} oraz F = p_0.",
            "Jeśli U ≤ F, zwrócić wartość 0.",
            "W przeciwnym razie obliczać p_k = p_{k-1}·λ/k.",
            "Zwrócić pierwsze k, dla którego U ≤ F.",
        ],
    )
    add_code(
        doc,
        """
def _poisson_from_uniform(uniform_value: float, lambda_value: float) -> tuple[int, float]:
    probability = exp(-lambda_value)
    cumulative = probability
    value = 0

    while uniform_value > cumulative:
        value += 1
        probability *= lambda_value / value
        cumulative += probability

    return value, cumulative
""",
    )
    add_paragraph(doc, "Aplikacja prezentuje wygenerowane wartości rozkładu Poissona, histogram, średnią z próby, wariancję z próby oraz porównanie z wartościami teoretycznymi.")

    doc.add_heading("Lista 2, zadanie 2", level=1)
    doc.add_heading("Treść zadania", level=2)
    add_paragraph(doc, "Zaimplementować metodologię metody odwracania dystrybuanty, wykorzystując wyniki Przykładu 1 z materiałów wykładowych.")
    add_paragraph(doc, "W projekcie jako przykład rozkładu ciągłego zastosowano rozkład wykładniczy z parametrem λ. Liczby jednostajne U_i generowane są za pomocą generatora LCG Lehmera.")
    doc.add_heading("Podstawy merytoryczne", level=2)
    add_equation(doc, "U ~ U(0,1),   X = F^{-1}(U)")
    add_equation(doc, "F(x)=0 dla x<0")
    add_equation(doc, "F(x)=1-e^{-λx} dla x≥0")
    add_equation(doc, "u = 1-e^{-λx}")
    add_equation(doc, "e^{-λx}=1-u")
    add_equation(doc, "x = -ln(1-u)/λ")
    add_equation(doc, "X = -ln(1-U)/λ")
    doc.add_heading("Algorytm", level=2)
    add_numbered(
        doc,
        [
            "Wygenerować wartość U algorytmem LCG Lehmera.",
            "Podstawić U do wzoru odwrotnej dystrybuanty.",
            "Otrzymać wartość X o rozkładzie wykładniczym.",
            "Powtórzyć procedurę dla zadanej liczby próbek.",
        ],
    )
    add_code(
        doc,
        """
uniform_value = min(max(generator.next_float(), 0.0), 0.999999999999)
value = -log(1 - uniform_value) / safe_lambda
""",
    )
    add_equation(doc, "E(X)=1/λ,   Var(X)=1/λ^2")

    doc.add_heading("Podsumowanie", level=1)
    add_paragraph(
        doc,
        "W projekcie opracowano trzy zadania z list laboratoryjnych. Każde zadanie zawiera "
        "podstawy teoretyczne, implementację algorytmu oraz prezentację wyników. Wspólnym "
        "elementem zadań symulacyjnych jest generator LCG Lehmera, który dostarcza liczby "
        "pseudolosowe z przedziału [0,1).",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
